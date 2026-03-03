"""
Tabular & Text Extractors — Real parsing for XLSX, CSV, DOCX, TXT files.
"""

import csv
import json
from io import StringIO

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def extract_xlsx(filepath):
    """Extract all sheets from an Excel file."""
    if not HAS_OPENPYXL:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}

    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    sheets = {}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        headers = []

        for i, row in enumerate(ws.iter_rows(values_only=True)):
            str_row = [str(c) if c is not None else "" for c in row]
            if i == 0:
                headers = str_row
            else:
                rows.append(str_row)

        sheets[sheet_name] = {
            "headers": headers,
            "rows": rows,
            "row_count": len(rows),
            "col_count": len(headers),
        }

    wb.close()

    # Primary sheet
    first_sheet = sheets[wb.sheetnames[0]] if wb.sheetnames else {"headers": [], "rows": [], "row_count": 0}

    return {
        "format": "xlsx",
        "sheets": sheets,
        "sheet_names": list(sheets.keys()),
        "primary_sheet": wb.sheetnames[0] if wb.sheetnames else None,
        "headers": first_sheet["headers"],
        "rows": first_sheet["rows"],
        "row_count": first_sheet["row_count"],
        "stats": {
            "total_sheets": len(sheets),
            "total_rows": sum(s["row_count"] for s in sheets.values()),
        },
    }


def extract_csv(filepath):
    """Extract data from a CSV file."""
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()

    # Auto-detect dialect
    try:
        dialect = csv.Sniffer().sniff(content[:4096])
    except csv.Error:
        dialect = csv.excel

    reader = csv.reader(StringIO(content), dialect)
    all_rows = list(reader)

    if not all_rows:
        return {"format": "csv", "headers": [], "rows": [], "row_count": 0, "stats": {"total_rows": 0}}

    headers = all_rows[0]
    rows = all_rows[1:]

    return {
        "format": "csv",
        "headers": headers,
        "rows": rows,
        "row_count": len(rows),
        "stats": {
            "total_rows": len(rows),
            "total_cols": len(headers),
        },
    }


def extract_docx(filepath):
    """Extract text and tables from a DOCX file."""
    if not HAS_DOCX:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    doc = DocxDocument(filepath)

    # Extract paragraphs
    paragraphs = []
    full_text = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append({
                "text": text,
                "style": para.style.name if para.style else "Normal",
            })
            full_text += text + "\n"

    # Extract tables
    tables = []
    for i, table in enumerate(doc.tables):
        rows = []
        headers = []
        for j, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if j == 0:
                headers = cells
            else:
                rows.append(cells)
        tables.append({
            "table_index": i,
            "headers": headers,
            "rows": rows,
            "row_count": len(rows),
        })

    return {
        "format": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
        "full_text": full_text,
        "stats": {
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables),
            "total_chars": len(full_text),
        },
    }


def extract_txt(filepath):
    """Extract text from a plain text file."""
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    return {
        "format": "txt",
        "full_text": text,
        "stats": {
            "total_chars": len(text),
            "total_lines": text.count("\n") + 1,
            "total_words": len(text.split()),
        },
    }
